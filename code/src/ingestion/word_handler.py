"""
Word document handler (.docx, .doc)
"""
from typing import Dict, Any, List
from pathlib import Path
from .base_handler import BaseHandler

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordHandler(BaseHandler):
    """Handler for Word documents"""
    
    def can_handle(self, file_path: str) -> bool:
        """Check if file is a Word document"""
        ext = Path(file_path).suffix.lower()
        return ext in ['.docx', '.doc']
    
    def get_supported_extensions(self) -> List[str]:
        return ['.docx', '.doc']
    
    def extract(self, source: str, **kwargs) -> Dict[str, Any]:
        """Extract data from Word document"""
        metadata = self.extract_metadata(source)
        result = {
            'data': [],
            'metadata': metadata,
            'text_content': [],
            'structure': {}
        }
        
        if not DOCX_AVAILABLE:
            raise ValueError(f"python-docx not installed. Install with: pip install python-docx")
        
        try:
            doc = Document(source)
            
            paragraphs_data = []
            all_text = []
            tables_data = []
            
            # Extract paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs_data.append({
                        'paragraph_number': para_idx + 1,
                        'text': text,
                        'style': para.style.name if para.style else None
                    })
                    all_text.append(text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = self._extract_table(table)
                tables_data.append({
                    'table_number': table_idx + 1,
                    'data': table_data
                })
            
            # Combine all text
            full_text = "\n\n".join(all_text)
            
            result['data'] = {
                'paragraphs': paragraphs_data,
                'tables': tables_data
            }
            
            result['text_content'] = [full_text]
            
            result['structure'] = {
                'format': 'word',
                'paragraph_count': len(paragraphs_data),
                'table_count': len(tables_data),
                'has_tables': len(tables_data) > 0
            }
            
            metadata.update({
                'paragraph_count': len(paragraphs_data),
                'table_count': len(tables_data)
            })
            
        except Exception as e:
            raise ValueError(f"Error reading Word document {source}: {str(e)}")
        
        return result
    
    def _extract_table(self, table) -> List[Dict]:
        """Extract data from a Word table"""
        table_data = []
        headers = []
        
        # First row as headers if available
        if len(table.rows) > 0:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # Extract rows
        start_row = 1 if headers else 0
        for row in table.rows[start_row:]:
            row_data = {}
            for col_idx, cell in enumerate(row.cells):
                header = headers[col_idx] if col_idx < len(headers) else f"Column_{col_idx}"
                row_data[header] = cell.text.strip()
            table_data.append(row_data)
        
        return table_data

